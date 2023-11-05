from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

document.add_picture(              #adding a picture
    'ahmed.jpg', 
     width=Inches(2.0)
)

#name phone and email details
speak('Hey buddy what is your name')
name = input('What is your name? ')
speak('Hello ' + name + ' How are you today?''hope you are good, enter your phone number please')

Phone_number = input('What is your Phone number? ')
speak('Great, buddy! now enter your email')
email = input('What is your email? ')

document.add_paragraph(
    name +' | ' + Phone_number + ' | ' + email)


document.add_heading('About Me')
speak('Now tell me about yourself')
document.add_paragraph(
    input('Tell me about yourself? ')
)

document.add_heading('Courses done')
speak('Very well! Now, Enter any courses if you have done')
p = document.add_paragraph(
    input('What courses have you done? ')
)
speak('Enter any skills that you have')
document.add_heading('Skills')
skill = input('Enter skills' )
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    skill = input("Do you want to add another skill or finish? ")
    if skill == "finish":
        break
    else:
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
        
speak('Thank you for the info ' + name + 'Have a great day!')
document.save('cv.docx')